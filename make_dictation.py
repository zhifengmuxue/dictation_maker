"""
把 Word (.docx) 文档中非黑色文字替换为等长的下划线式空格（用于默写版）。

用法:
    python make_dictation.py input.docx

会在同目录输出 `input_dictation.docx`。

实现说明：
 - 识别 run.font.color.rgb 存在且 != RGBColor(0x00,0x00,0x00) 视为非黑色。
 - 替换文本为每个原字符对应两个非断行空格（\u00A0），并设置 run.font.underline = True，颜色改为黑色。
 - 遍历段落和表格单元格。可扩展到页眉/页脚或其他 color 表示方式。
"""
import sys
from docx import Document
from docx.shared import RGBColor


def is_non_black(run):
    """判断 run 是否为非黑色文本（有明确 RGB 且不等于黑色）。"""
    color = run.font.color
    if color is None:
        # 没有显式颜色（通常为自动/默认），视为黑色，不替换
        return False
    rgb = color.rgb
    if rgb is None:
        # 颜色存在但没有 RGB 信息（例如 theme），保守处理为黑色，不替换
        return False
    # 如果明确的 RGB 不等于黑色，则视为非黑色
    return rgb != RGBColor(0x00, 0x00, 0x00)


def replace_non_black_in_run(run):
    text = run.text
    if not text:
        return
    if is_non_black(run):
        length = len(text)
        # 每个字符对应两个非断行空格，使下划线更长
        under_text = '\u00A0' * (length * 2)

        # 先尝试清除 theme_color（如果存在），并预先设置为黑色与下划线，
        # 然后再替换文本，最后再次确保属性已应用。
        try:
            if run.font.color is not None and hasattr(run.font.color, 'theme_color'):
                run.font.color.theme_color = None
        except Exception:
            # 如果无法清除 theme_color，继续但不终止
            pass

        # 先设置样式
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.underline = True

        # 替换文本
        run.text = under_text

        # 再次确保样式已应用（有时赋值文本会导致样式丢失）
        run.font.underline = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def process_paragraph(paragraph):
    for run in paragraph.runs:
        replace_non_black_in_run(run)


def process_table(table):
    for row in table.rows:
        for cell in row.cells:
            process_block(cell)


def process_block(block):
    # block may be Document, _Cell, Header, Footer ... they have .paragraphs and .tables
    for p in getattr(block, 'paragraphs', []):
        process_paragraph(p)
    for t in getattr(block, 'tables', []):
        process_table(t)


def make_dictation(in_path, out_path=None):
    doc = Document(in_path)
    process_block(doc)
    # TODO: 可扩展处理 headers/footers
    if out_path is None:
        if in_path.lower().endswith('.docx'):
            out_path = in_path[:-5] + '_dictation.docx'
        else:
            out_path = in_path + '_dictation.docx'
    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 2:
        print('用法: python make_dictation.py input.docx')
        sys.exit(1)
    in_path = sys.argv[1]
    out = make_dictation(in_path)
    print('输出：', out)


if __name__ == '__main__':
    main()
